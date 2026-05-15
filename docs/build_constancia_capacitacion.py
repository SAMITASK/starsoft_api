from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path(__file__).resolve().parent / "constancia-capacitacion-starcheck.docx"

PARTICIPANTES = [
    "BUSTAMANTE PAZ YVONNE SALLI",
    "CUBAS RAMOS HIPOLITO",
    "CHAHUA MACCAPA LUIS MIGUEL",
]


def set_cell_border(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "AEB7C6")


def set_row_height(row, height_cm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def style_document(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)


def add_header(document: Document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("CONSTANCIA DE CAPACITACION")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistema StarCheck")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = None


def add_body(document: Document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.35

    text = (
        "Por medio del presente se deja constancia de que se realizo la capacitacion "
        "sobre el uso del sistema StarCheck a las siguientes personas, quienes firman "
        "en senal de participacion y conformidad."
    )
    paragraph.add_run(text)

    info = document.add_paragraph()
    info.paragraph_format.space_after = Pt(12)
    info.paragraph_format.line_spacing = 1.2
    info.add_run("Fecha de capacitacion: ").bold = True
    info.add_run("15 de mayo de 2026")


def add_signature_table(document: Document):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Cm(1.4), Cm(10.8), Cm(5.4)]
    headers = ["N°", "Participante", "Firma"]

    for index, cell in enumerate(table.rows[0].cells):
        cell.width = widths[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(headers[index])
        run.bold = True
        run.font.size = Pt(11)
        set_cell_border(cell)

    for index, participante in enumerate(PARTICIPANTES, start=1):
        row = table.add_row()
        set_row_height(row, 2.4)

        number_cell, name_cell, sign_cell = row.cells
        number_cell.width = widths[0]
        name_cell.width = widths[1]
        sign_cell.width = widths[2]

        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)

        number_paragraph = number_cell.paragraphs[0]
        number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number_paragraph.add_run(str(index)).font.size = Pt(11)

        name_paragraph = name_cell.paragraphs[0]
        name_paragraph.paragraph_format.space_before = Pt(8)
        name_paragraph.paragraph_format.space_after = Pt(8)
        run = name_paragraph.add_run(participante)
        run.font.size = Pt(11)

        sign_paragraph = sign_cell.paragraphs[0]
        sign_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sign_paragraph.paragraph_format.space_before = Pt(34)
        sign_paragraph.add_run("______________________").font.size = Pt(11)


def add_footer_blocks(document: Document):
    document.add_paragraph("")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(18)
    note.paragraph_format.line_spacing = 1.2
    note.add_run("Observacion: ").bold = True
    note.add_run("Documento preparado para firma de los participantes capacitados.")

    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing.paragraph_format.space_after = Pt(24)
    closing.add_run("Lima, 15 de mayo de 2026")

    trainer = document.add_paragraph()
    trainer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trainer.paragraph_format.space_before = Pt(24)
    trainer.add_run("__________________________________")

    trainer_name = document.add_paragraph()
    trainer_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trainer_name.paragraph_format.space_before = Pt(2)
    trainer_name.add_run("Responsable de la capacitacion").bold = True


def main():
    document = Document()
    style_document(document)
    add_header(document)
    add_body(document)
    add_signature_table(document)
    add_footer_blocks(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
